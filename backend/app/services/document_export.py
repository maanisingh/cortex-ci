"""
Document Export Service for Russian Compliance Documents
Supports DOCX and PDF export formats
"""

import io
from datetime import date
from typing import Any
from uuid import UUID

import structlog

logger = structlog.get_logger()


class DocumentExportService:
    """
    Service for exporting compliance documents to DOCX and PDF formats.
    """

    def __init__(self):
        self._docx_available = False
        self._pdf_available = False
        self._check_dependencies()

    def _check_dependencies(self):
        """Check if export dependencies are available."""
        try:
            from docx import Document
            self._docx_available = True
        except ImportError:
            logger.warning("python-docx not installed, DOCX export disabled")

        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.pdfgen import canvas
            self._pdf_available = True
        except ImportError:
            logger.warning("reportlab not installed, PDF export disabled")

    async def export_to_docx(
        self,
        content: str,
        title: str,
        company_name: str | None = None,
        metadata: dict[str, Any] | None = None,
    ) -> bytes:
        """
        Export document content to DOCX format.

        Args:
            content: Document text content
            title: Document title
            company_name: Company name for header
            metadata: Additional document metadata

        Returns:
            DOCX file as bytes
        """
        if not self._docx_available:
            raise RuntimeError("python-docx is not installed")

        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Set document properties
        core_properties = doc.core_properties
        core_properties.title = title
        if company_name:
            core_properties.author = company_name
        core_properties.created = date.today()

        # Add header with company name
        if company_name:
            header = doc.sections[0].header
            header_para = header.paragraphs[0]
            header_para.text = company_name
            header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Add title
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add content - split by newlines and process
        lines = content.split('\n')
        for line in lines:
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph()
            elif stripped.startswith('#'):
                # Markdown-style heading
                level = len(stripped) - len(stripped.lstrip('#'))
                heading_text = stripped.lstrip('# ')
                doc.add_heading(heading_text, level=min(level, 9))
            elif stripped.isupper() and len(stripped) < 100:
                # All-caps section header
                para = doc.add_paragraph()
                run = para.add_run(stripped)
                run.bold = True
            else:
                doc.add_paragraph(stripped)

        # Add footer with page number
        footer = doc.sections[0].footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"Generated: {date.today().strftime('%d.%m.%Y')}"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    async def export_to_pdf(
        self,
        content: str,
        title: str,
        company_name: str | None = None,
        metadata: dict[str, Any] | None = None,
    ) -> bytes:
        """
        Export document content to PDF format.

        Args:
            content: Document text content
            title: Document title
            company_name: Company name for header
            metadata: Additional document metadata

        Returns:
            PDF file as bytes
        """
        if not self._pdf_available:
            raise RuntimeError("reportlab is not installed")

        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        buffer = io.BytesIO()

        # Create document
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=2*cm,
            leftMargin=2*cm,
            topMargin=2*cm,
            bottomMargin=2*cm,
        )

        # Get styles
        styles = getSampleStyleSheet()

        # Create custom styles for Russian text
        title_style = ParagraphStyle(
            'RuTitle',
            parent=styles['Heading1'],
            fontSize=16,
            alignment=1,  # Center
            spaceAfter=20,
        )

        heading_style = ParagraphStyle(
            'RuHeading',
            parent=styles['Heading2'],
            fontSize=12,
            spaceBefore=12,
            spaceAfter=6,
        )

        body_style = ParagraphStyle(
            'RuBody',
            parent=styles['Normal'],
            fontSize=10,
            leading=14,
            spaceBefore=6,
            spaceAfter=6,
        )

        # Build content
        story = []

        # Add company name header
        if company_name:
            story.append(Paragraph(company_name, styles['Normal']))
            story.append(Spacer(1, 0.5*cm))

        # Add title
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 0.5*cm))

        # Process content
        lines = content.split('\n')
        for line in lines:
            stripped = line.strip()
            if not stripped:
                story.append(Spacer(1, 0.3*cm))
            elif stripped.isupper() and len(stripped) < 100:
                story.append(Paragraph(f"<b>{stripped}</b>", heading_style))
            else:
                # Escape special characters for reportlab
                escaped = stripped.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                story.append(Paragraph(escaped, body_style))

        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()

    @property
    def docx_available(self) -> bool:
        """Check if DOCX export is available."""
        return self._docx_available

    @property
    def pdf_available(self) -> bool:
        """Check if PDF export is available."""
        return self._pdf_available


# Singleton instance
document_export_service = DocumentExportService()
